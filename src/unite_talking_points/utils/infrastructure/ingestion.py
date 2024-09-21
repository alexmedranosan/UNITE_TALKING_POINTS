from typing import List

import PyPDF2
import docx

from src.unite_talking_points.domain.entities.entities import Document
from src.unite_talking_points.utils.directory.directory_utils import get_file_paths


def load_pdf_document(path: str) -> Document:
    """
    Load a PDF document and transform it into a Document object.
    :param path: str Path to the PDF file.
    :return: document: Document object
    """
    with open(path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        num_pages = len(reader.pages)

        text = ''
        for page_number in range(num_pages):
            page = reader.pages[page_number]
            text += page.extract_text()

        # Extract metadata
        author = reader.metadata.get("Author", None)
        date_created = reader.metadata.get("creation_date", None)
        date_modified = reader.metadata.get("modification_date", None)

    # Create the Document object
    document = Document(content=text,
                        _id=None,
                        origin=None,
                        title=None,
                        author=author,
                        keywords=None,
                        date_created=date_created,
                        date_modified=date_modified,
                        source=path)

    return document


def load_word_document(path: str) -> Document:
    """
    Load a Word document and transform it into a Document object.
    :param path: str Path to the Word document.
    :return: document: Document object.
    """
    doc = docx.Document(path)

    text = []
    for page in doc.paragraphs:
        text.append(page.text)

    # Extract the metadata from the document
    author = doc.core_properties.author
    date_created = doc.core_properties.created
    date_modified = doc.core_properties.modified

    # Create the Document object
    document = Document(content='\n'.join(text),
                        _id=None,
                        origin=None,
                        title=None,
                        author=author,
                        keywords=None,
                        date_created=date_created,
                        date_modified=date_modified,
                        source=path)

    return document


def load_documents(directory: str) -> List[Document]:
    """
    Loads all documents in a given directory into a list of Documents objects.
    :param directory: str The directory path.
    :return: documents: List[Document] A list of Documents objects.
    """
    file_list = get_file_paths(directory)
    documents = []

    for path in file_list:
        if path.endswith(".pdf") or path.endswith(".PDF"):
            document = load_pdf_document(path)

        elif path.endswith(".docx"):
            document = load_word_document(path)

        else:
            print(f"WARNING: Unsupported document extension for: {path}")
            continue

        if document.content:
            documents.append(document)

        else:
            print(f"WARNING: Empty document found for: {path}")

    return documents
